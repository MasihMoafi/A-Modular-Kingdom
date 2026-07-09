#!/usr/bin/env python
# coding: utf-8
import os
os.environ["TQDM_DISABLE"] = "1"  # Disable tqdm via env var
# Prefer offline-first HF/Transformers to avoid tool timeouts in restricted envs.
# Users can override these before starting the server.
os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
os.environ.setdefault("HF_HUB_DISABLE_TELEMETRY", "1")
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

import sys
import io

repo_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
project_root = os.path.join(repo_root, "src")
workspace_root = os.path.realpath(os.environ.get("AMK_WORKSPACE_ROOT", os.getcwd()))
sys.path.insert(0, project_root)
import json
import importlib
import importlib.util
import glob
import shutil
import subprocess
import zipfile
from xml.sax.saxutils import escape as _xml_escape
from datetime import datetime
from pathlib import Path
from typing import Dict, List
from mcp.server.fastmcp import FastMCP
from pydantic import Field
from memory.scoped_manager import ScopedMemoryManager
from memory.memory_config import MemoryScope
from memory.path_policy import resolve_memory_base
from agent.tool_execution import call_in_subprocess as _call_in_subprocess_impl
from agent.tool_execution import call_tool_safely as _call_tool_safely_impl

# Default MCP surface area: RAG + scoped memory.
# Extra tools are optional behind env flags to keep startup/tooling lean.
_EXPOSE_EXTRA_TOOLS = os.environ.get("MCP_EXPOSE_EXTRA_TOOLS", "1").lower() in ("1", "true", "yes", "y", "on")
_EXPOSE_RESOURCES = os.environ.get("MCP_EXPOSE_RESOURCES", "1").lower() in ("1", "true", "yes", "y", "on")

_DOC_EXTENSIONS = {
    ".md", ".markdown", ".txt", ".py", ".json", ".yaml", ".yml", ".toml",
    ".html", ".css", ".js", ".ts", ".tsx", ".jsx", ".ipynb", ".pdf", ".docx",
}
_DOC_SKIP_DIRS = {
    ".git", ".venv", "__pycache__", ".pytest_cache", "agent_chroma_db",
    "agent_qdrant_db", "rag_db_v1", "rag_db_v2", "qdrant_storage",
}

_DEFAULT_MEMORY_BASE = resolve_memory_base(project_root=workspace_root)
os.environ.setdefault("MEMORY_BASE_PATH", _DEFAULT_MEMORY_BASE)

# Initialize FastMCP
mcp = FastMCP("unified_knowledge_agent_host")
scoped_memory = None

_LOG_STDERR = os.environ.get("MCP_LOG_STDERR", "0").lower() in ("1", "true", "yes", "y", "on")
_LOG_FILE = os.environ.get("MCP_LOG_FILE", "").strip()
_DEBUG_PROTOCOL = os.environ.get("MCP_DEBUG_PROTOCOL", "0").lower() in ("1", "true", "yes", "y", "on")


def _memory_counts(manager: ScopedMemoryManager) -> dict[str, int]:
    counts = {}
    for scope in MemoryScope:
        try:
            counts[scope.value] = len(manager.list_all(scope))
        except Exception:
            counts[scope.value] = 0
    return counts


def _legacy_project_roots() -> list[str]:
    candidates = [project_root]
    roots: list[str] = []
    seen = set()
    for root in candidates:
        rr = os.path.realpath(root)
        if rr in seen:
            continue
        seen.add(rr)
        legacy_dir = Path(rr) / ".modular_kingdom" / "memories"
        if legacy_dir.exists():
            roots.append(rr)
    return roots


def _legacy_extra_dirs() -> list[str]:
    return [
        str(Path(project_root) / ".modular_kingdom"),
        str(Path(project_root) / "agent_chroma_db"),
        str(Path(project_root) / "data" / "agent_chroma_db"),
        str(Path(project_root) / "agent_qdrant_db"),
    ]


def _scoped_manager_for_root(project_root_value: str, force_project_local: bool = False) -> ScopedMemoryManager:
    if not force_project_local:
        return ScopedMemoryManager(project_root=project_root_value)

    previous_memory_base = os.environ.pop("MEMORY_BASE_PATH", None)
    previous_backend = os.environ.get("AMK_MEMORY_VECTOR_BACKEND")
    os.environ["AMK_MEMORY_VECTOR_BACKEND"] = "qdrant"
    try:
        return ScopedMemoryManager(project_root=project_root_value)
    finally:
        if previous_memory_base is not None:
            os.environ["MEMORY_BASE_PATH"] = previous_memory_base
        if previous_backend is None:
            os.environ.pop("AMK_MEMORY_VECTOR_BACKEND", None)
        else:
            os.environ["AMK_MEMORY_VECTOR_BACKEND"] = previous_backend


def _resolve_target_path(path_str: str) -> Path:
    p = Path(path_str.strip()).expanduser()
    if not p.is_absolute():
        p = Path(workspace_root) / p
    return p.resolve()


def _is_safe_target(path: Path) -> bool:
    allowed = [Path(workspace_root).resolve(), Path(repo_root).resolve(), Path("/tmp").resolve()]
    rp = path.resolve()
    for root in allowed:
        if rp == root or str(rp).startswith(str(root) + os.sep):
            return True
    return False


def _write_simple_docx(path: Path, title: str, body: str) -> None:
    title = (title or "").strip()
    body = (body or "").replace("\r\n", "\n").strip()
    paragraphs = []
    if title:
        paragraphs.append(f'<w:p><w:r><w:t>{_xml_escape(title)}</w:t></w:r></w:p>')
    if body:
        for line in body.split("\n"):
            txt = _xml_escape(line) if line else ""
            if txt:
                paragraphs.append(f'<w:p><w:r><w:t xml:space="preserve">{txt}</w:t></w:r></w:p>')
            else:
                paragraphs.append("<w:p/>")
    if not paragraphs:
        paragraphs.append("<w:p/>")

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'mc:Ignorable="w14 w15 wp14">'
        "<w:body>"
        + "".join(paragraphs)
        + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" '
        'w:bottom="1440" w:left="1440" w:header="708" w:footer="708" w:gutter="0"/>'
        '<w:cols w:space="708"/><w:docGrid w:linePitch="360"/></w:sectPr>'
        "</w:body></w:document>"
    )

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )

    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)


def _elog(msg: str) -> None:
    try:
        if _LOG_STDERR:
            sys.stderr.write(msg)
            sys.stderr.flush()
        if _LOG_FILE:
            with open(_LOG_FILE, "a", encoding="utf-8") as f:
                f.write(msg)
    except Exception:
        pass


def _call_in_subprocess(module_name: str, function_name: str, payload: dict, timeout: int = 120) -> str:
    return _call_in_subprocess_impl(
        module_name=module_name,
        function_name=function_name,
        payload=payload,
        project_root=project_root,
        timeout=timeout,
    )


def _call_tool_safely(func, *args, **kwargs):
    return _call_tool_safely_impl(func, _elog, *args, **kwargs)


try:
    # Clear stale Qdrant lock files from previous crashes
    import glob as _glob
    _mem_root = os.path.join(workspace_root, ".modular_kingdom")
    for _lock in _glob.glob(os.path.join(_mem_root, "**", ".lock"), recursive=True):
        os.remove(_lock)
        _elog(f"[HOST] Removed stale lock: {_lock}\n")

    scoped_memory = ScopedMemoryManager(project_root=workspace_root)
    _elog("[HOST] Scoped memory initialized successfully\n")
except Exception as e:
    _elog(f"[HOST] WARNING: Memory system disabled: {e}\n")
    scoped_memory = None

# Restore HTTP(S) proxy for RAG subprocesses (SOCKS breaks httpx)
from utils.proxy import restore_http as _restore_proxy
_restore_proxy()

@mcp.tool(
    name="save_fact",
    description="Save structured facts to the memory system with automatic processing and fact extraction from the provided content"
)
def save_fact(
    fact_data: Dict = Field(description="Dictionary containing 'content' key with the fact to save")
) -> str:
    if scoped_memory is None:
        return json.dumps({"error": "Scoped memory not initialized."})
    try:
        content_to_save = fact_data.get('content')
        if not content_to_save:
            return json.dumps({"error": "No 'content' field in fact JSON."})
        memory_id = scoped_memory.save(content_to_save)
        return json.dumps({"status": "success", "id": memory_id, "message": f"Fact saved: '{content_to_save}'"})
    except Exception as e:
        return json.dumps({"error": f"Error saving fact: {str(e)}"})

@mcp.tool(
    name="save_memory",
    description="Persist a user-approved memory. Use for explicit save requests only; supports global, persona, and project scopes."
)
def save_direct_memory(
    content: str = Field(description="Text to save (with optional #scope:type prefix)"),
    scope: str = Field(default="", description="Explicit scope: 'global_rules', 'global_preferences', 'project_context', etc.")
) -> str:
    """Save content to scoped memory."""
    if scoped_memory is None:
        return json.dumps({"error": "Scoped memory not initialized."})
    try:
        # Parse scope from content prefix if present
        target_scope = None
        if scope:
            try:
                target_scope = MemoryScope(scope)
            except ValueError:
                pass
        
        # Check for prefix in content
        if not target_scope and content.startswith("#"):
            from memory.memory_config import MemoryConfig
            config = MemoryConfig()
            parsed_scope, cleaned_content = config.parse_scope_prefix(content)
            if parsed_scope:
                target_scope = parsed_scope
                content = cleaned_content
        
        memory_id = scoped_memory.save(content, scope=target_scope)
        scope_name = target_scope.value if target_scope else "inferred"
        return json.dumps({"status": "success", "scope": scope_name, "id": memory_id})
    except Exception as e:
        return json.dumps({"error": f"Error saving scoped memory: {str(e)}"})

@mcp.tool(
    name="set_global_rule",
    description="Set a permanent global rule that persists across all projects and sessions"
)
def set_global_rule(
    rule: str = Field(description="The rule/instruction to save globally")
) -> str:
    """Convenience tool to save global rules."""
    if scoped_memory is None:
        return json.dumps({"error": "Scoped memory not initialized."})
    try:
        memory_id = scoped_memory.save(rule, scope=MemoryScope.GLOBAL_RULES)
        return json.dumps({"status": "success", "scope": "global_rules", "id": memory_id})
    except Exception as e:
        return json.dumps({"error": f"Error setting global rule: {str(e)}"})

@mcp.tool(
    name="delete_memory",
    description="Delete a specific memory from the memory system using its unique identifier"
)
def delete_memory(
    memory_id: str = Field(description="The unique ID of the memory to delete")
) -> str:
    """Delete a memory by its ID (searches all scopes)."""
    if scoped_memory is None:
        return json.dumps({"error": "Scoped memory not initialized."})
    try:
        for scope in MemoryScope:
            try:
                scoped_memory.delete(memory_id, scope)
                return json.dumps({"status": "success", "message": f"Memory {memory_id} deleted from {scope.value}"})
            except KeyError:
                continue  # Not found in this scope
            except Exception as e:
                sys.stderr.write(f"[HOST] Delete error in {scope.value}: {e}\n")
                continue  # Log but try next scope
        return json.dumps({"error": f"Memory {memory_id} not found in any scope"})
    except Exception as e:
        return json.dumps({"error": f"Error deleting memory: {str(e)}"})

@mcp.tool(
    name="search_memories",
    description="Search durable scoped memory for relevant rules, preferences, persona notes, and project context. Returns content plus scope/source metadata."
)
def search_memories(
    query: str = Field(description="Search query"),
    top_k: int = Field(default=3, description="Results per scope"),
    scope_filter: str = Field(default="", description="Optional: search only specific scope")
) -> str:
    if scoped_memory is None:
        return json.dumps([{"error": "Scoped memory not initialized."}])
    try:
        scopes_to_search = None
        if scope_filter:
            try:
                scopes_to_search = [MemoryScope(scope_filter)]
            except ValueError:
                pass
        
        results = scoped_memory.search(query, k=top_k, scopes=scopes_to_search)
        formatted = []
        for res in results:
            metadata = res.get("metadata", {})
            formatted.append({
                "content": res["content"],
                "scope": metadata.get("scope", "unknown"),
                "id": res["id"],
                "source": metadata.get("source"),
                "start_line": metadata.get("start_line"),
                "end_line": metadata.get("end_line"),
                "score": metadata.get("score"),
                "text_score": metadata.get("text_score"),
                "vector_score": metadata.get("vector_score"),
            })
        return json.dumps(formatted)
    except Exception as e:
        return json.dumps([{"error": f"Error searching scoped memories: {str(e)}"}])

@mcp.tool(
    name="query_knowledge_base",
    description="Run local RAG over the launch workspace or supplied files/dir. Use for broad exploration; read exact files with read_file when the user names them."
)
def query_knowledge_base(
    query: str = Field(description="The search query for the knowledge base"),
    version: str = Field(default="v1", description="RAG version to use: 'v1' or 'v2'"),
    doc_path: str = Field(default="", description="Optional path to a specific documents directory"),
    files: List[str] = Field(default=[], description="Optional list of specific file paths to index and search")
) -> str:
    """
    A tool to query the external knowledge base (RAG system) with selectable versions.
    """
    _elog(f"[HOST] RAG tool called with query: '{query[:20]}...', version: {version}, path: '{doc_path}', files: {len(files) if files else 0}\n")

    try:
        # Smart defaults:
        # - If neither `doc_path` nor `files` were provided, index the launch
        #   workspace by default.
        # - Treat relative paths as relative to that launch workspace.
        repo_default = workspace_root
        shortcut_words = {"desktop", "documents", "downloads"}

        if (not doc_path) and (not files):
            doc_path = repo_default

        if doc_path and isinstance(doc_path, str):
            dp = doc_path.strip()
            if dp and (not os.path.isabs(dp)) and (not dp.startswith("~")) and (dp.lower() not in shortcut_words):
                doc_path = os.path.join(repo_default, dp)

        if files:
            normalized = []
            for f in files:
                if not f or not isinstance(f, str):
                    continue
                fp = f.strip()
                if not fp:
                    continue
                if (not os.path.isabs(fp)) and (not fp.startswith("~")):
                    fp = os.path.join(repo_default, fp)
                normalized.append(fp)
            files = normalized

        # Map version to module name
        module_map = {
            "v1": "rag.fetch_1",
            "v2": "rag.fetch_2"
        }

        module_name = module_map.get(version)
        if not module_name:
            return json.dumps({"error": f"Invalid RAG version '{version}'. Must be 'v1' or 'v2'."})

        # IMPORTANT: Keep RAG in-process so models/pipelines can be cached across calls.
        # The previous per-call subprocess approach was reliable for isolation, but it
        # makes cold-start latency huge and frequently exceeds client tool timeouts.
        _elog(f"[HOST] Calling fetchExternalKnowledge from {module_name} in-process\n")

        payload = {"query": query}
        if files and len(files) > 0:
            payload["file_list"] = files
        elif doc_path and isinstance(doc_path, str):
            payload["doc_path"] = doc_path

        module = importlib.import_module(module_name)
        func = getattr(module, "fetchExternalKnowledge", None)
        if not callable(func):
            return json.dumps({"error": f"Module '{module_name}' has no callable fetchExternalKnowledge(query, ...)"})

        def _run_rag():
            # Keep signature explicit to avoid silent arg mismatches.
            if "file_list" in payload:
                return func(query=payload["query"], file_list=payload["file_list"])
            if "doc_path" in payload:
                return func(query=payload["query"], doc_path=payload["doc_path"])
            return func(query=payload["query"])

        results = _call_tool_safely(_run_rag)

        _elog(f"[HOST] RAG ({version}) returned {len(str(results)) if results else 0} chars\n")

        return json.dumps({"result": results})

    except (ModuleNotFoundError, AttributeError) as e:
        error_msg = f"Could not find or load 'fetchExternalKnowledge' function in module for version '{version}': {str(e)}"
        _elog(f"[HOST] RAG error: {error_msg}\n")
        return json.dumps({"error": error_msg})
    except Exception as e:
        import traceback
        error_msg = f"Error querying knowledge base on host with version {version}: {str(e)}"
        _elog(f"[HOST] RAG error: {error_msg}\n")
        _elog(f"[HOST] Traceback: {traceback.format_exc()}\n")
        return json.dumps({"error": error_msg})

@mcp.tool(
    name="list_all_memories",
    description="Retrieve and list all saved memories from the memory system for review and management"
)
def list_all_memories() -> str:
    """Retrieve and list all saved memories from the memory system for review and management"""
    _elog("[HOST] list_all_memories called\n")
    if not scoped_memory:
        return json.dumps({"error": "Memory system not initialized"})
    
    try:
        all_memories = {}
        
        for scope in MemoryScope:
            try:
                memories = scoped_memory.list_all(scope)
                all_memories[scope.value] = memories
            except Exception as e:
                _elog(f"[HOST] Error listing memories for scope {scope}: {e}\n")
        
        return json.dumps(all_memories, indent=2)
    except Exception as e:
        _elog(f"[HOST] Critical error in list_all_memories: {e}\n")
        import traceback
        traceback.print_exc()
        return json.dumps({"error": str(e)})

@mcp.tool(
    name="health_check",
    description="Check MCP server health and component status"
)
def health_check() -> str:
    """Return status of all MCP components."""
    status = {
        "server": "ok",
        "memory": "ok" if scoped_memory else "disabled",
    }
    # Keep health checks cheap. Importing rag.fetch_1 pulls heavier retrieval
    # dependencies and made simple health probes take several seconds.
    try:
        status["rag"] = "available" if importlib.util.find_spec("rag.fetch_1") else "missing"
    except Exception as e:
        status["rag"] = f"error: {str(e)}"
    return json.dumps(status)


@mcp.tool(
    name="web_search",
    description="Search the public web for current or unavailable local information. Prefer local files/memory first unless the user asks for web/current info."
)
def web_search(
    query: str = Field(description="The search query")
) -> str:
    """Always-available web search tool."""
    try:
        return _call_in_subprocess("tools.web_search", "perform_web_search", {"query": query}, timeout=60)
    except Exception as e:
        return json.dumps({"error": f"Web search failed: {str(e)}"})


@mcp.tool(
    name="crawl_web",
    description="Fetch a small set of web pages, convert them to local markdown, and return files for downstream RAG. Use after web_search when page content is needed."
)
def crawl_web(
    query: str = Field(default="", description="Search query used to find seed pages"),
    urls: List[str] = Field(default=[], description="Explicit seed URLs (optional)"),
    max_pages: int = Field(default=5, description="Maximum pages to crawl"),
    max_depth: int = Field(default=1, description="Link-follow depth from seed pages"),
    same_domain_only: bool = Field(default=True, description="Follow only same-domain links when crawling"),
    output_dir: str = Field(default="/tmp/web_crawl", description="Directory for crawled markdown files")
) -> str:
    """Crawl and normalize web content into local files."""
    try:
        payload = {
            "query": query,
            "urls": urls,
            "max_pages": max_pages,
            "max_depth": max_depth,
            "same_domain_only": same_domain_only,
            "output_dir": output_dir,
        }
        return _call_in_subprocess("tools.web_crawler", "crawl_webpages", payload, timeout=240)
    except Exception as e:
        return json.dumps({"error": f"crawl_web failed: {str(e)}"})


@mcp.tool(
    name="read_file",
    description="Read a local file or list a directory. Relative paths resolve under AMK_WORKSPACE_ROOT, the directory where the user launched the agent."
)
def read_file(
    path: str = Field(description="Absolute or relative file path"),
    max_chars: int = Field(default=12000, description="Max chars to return")
) -> str:
    try:
        target = _resolve_target_path(path)
        if not _is_safe_target(target):
            return json.dumps({"error": "Path blocked. Allowed roots: workspace, AMK repo, and /tmp."})
        if target.is_dir():
            entries = []
            for child in sorted(target.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower())):
                entries.append(child.name + ("/" if child.is_dir() else ""))
            return json.dumps(
                {
                    "status": "success",
                    "path": str(target),
                    "content": "\n".join(entries) if entries else "(empty directory)",
                    "is_directory": True,
                    "truncated": False,
                }
            )
        if not target.is_file():
            return json.dumps({"error": f"Not a file: {str(target)}"})
        text = target.read_text(encoding="utf-8", errors="replace")
        truncated = False
        if len(text) > max_chars:
            text = text[:max_chars]
            truncated = True
        return json.dumps(
            {
                "status": "success",
                "path": str(target),
                "content": text,
                "truncated": truncated,
            }
        )
    except Exception as e:
        return json.dumps({"error": f"read_file failed: {str(e)}"})


@mcp.tool(
    name="write_file",
    description="Create or overwrite a local text file. Relative paths resolve under AMK_WORKSPACE_ROOT; writes are limited to allowed local roots."
)
def write_file(
    path: str = Field(description="Absolute or relative file path"),
    content: str = Field(description="File content"),
    overwrite: bool = Field(default=True, description="Allow overwrite when file exists")
) -> str:
    try:
        target = _resolve_target_path(path)
        if not _is_safe_target(target):
            return json.dumps({"error": "Path blocked. Allowed roots: workspace, AMK repo, and /tmp."})
        if target.exists() and not overwrite:
            return json.dumps({"error": f"File exists and overwrite is false: {str(target)}"})
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content or "", encoding="utf-8")
        return json.dumps(
            {
                "status": "success",
                "path": str(target),
                "bytes": len((content or "").encode("utf-8")),
            }
        )
    except Exception as e:
        return json.dumps({"error": f"write_file failed: {str(e)}"})


@mcp.tool(
    name="create_markdown",
    description="Create a Markdown file (.md)."
)
def create_markdown(
    path: str = Field(description="Target .md path"),
    title: str = Field(default="", description="Optional title for H1"),
    body: str = Field(default="", description="Markdown body"),
    overwrite: bool = Field(default=True, description="Allow overwrite")
) -> str:
    try:
        target = _resolve_target_path(path)
        if target.suffix.lower() not in (".md", ".markdown"):
            return json.dumps({"error": "Path must end with .md or .markdown"})
        parts = []
        if title.strip():
            parts.append(f"# {title.strip()}")
            parts.append("")
        if body:
            parts.append(body)
        content = "\n".join(parts).strip() + "\n"
        return write_file(path=str(target), content=content, overwrite=overwrite)
    except Exception as e:
        return json.dumps({"error": f"create_markdown failed: {str(e)}"})


@mcp.tool(
    name="create_docx",
    description="Create a .docx document without external APIs."
)
def create_docx(
    path: str = Field(description="Target .docx path"),
    title: str = Field(default="", description="Optional document title"),
    body: str = Field(default="", description="Document body text"),
    overwrite: bool = Field(default=True, description="Allow overwrite")
) -> str:
    try:
        target = _resolve_target_path(path)
        if target.suffix.lower() != ".docx":
            return json.dumps({"error": "Path must end with .docx"})
        if not _is_safe_target(target):
            return json.dumps({"error": "Path blocked. Allowed roots: workspace, AMK repo, and /tmp."})
        if target.exists() and not overwrite:
            return json.dumps({"error": f"File exists and overwrite is false: {str(target)}"})
        try:
            from docx import Document  # type: ignore

            doc = Document()
            if (title or "").strip():
                doc.add_heading((title or "").strip(), level=1)
            text = (body or "").replace("\r\n", "\n")
            if text.strip():
                blocks = [b for b in text.split("\n\n")]
                for block in blocks:
                    doc.add_paragraph(block)
            doc.save(str(target))
        except Exception:
            # Fallback for offline/minimal environments where python-docx is unavailable.
            _write_simple_docx(target, title=title, body=body)
        return json.dumps(
            {
                "status": "success",
                "path": str(target),
                "bytes": target.stat().st_size if target.exists() else 0,
            }
        )
    except Exception as e:
        return json.dumps({"error": f"create_docx failed: {str(e)}"})


@mcp.tool(
    name="memory_storage_info",
    description="Show active memory path, per-scope counts, and detected legacy memory locations."
)
def memory_storage_info() -> str:
    if not scoped_memory:
        return json.dumps({"error": "Memory system not initialized"})
    try:
        active_root = os.environ.get("MEMORY_BASE_PATH", _DEFAULT_MEMORY_BASE)
        legacy_roots = _legacy_project_roots()
        legacy_details = []
        for root in legacy_roots:
            try:
                mgr = _scoped_manager_for_root(root, force_project_local=True)
                legacy_details.append({
                    "project_root": root,
                    "counts": _memory_counts(mgr),
                })
            except Exception as e:
                legacy_details.append({"project_root": root, "error": str(e)})

        extra_legacy_dirs = _legacy_extra_dirs()

        return json.dumps(
            {
                "active_memory_base": active_root,
                "active_project_root": workspace_root,
                "active_backend": "markdown",
                "active_counts": _memory_counts(scoped_memory),
                "active_markdown_files": scoped_memory.markdown_store.storage_files(),
                "legacy_project_roots": legacy_details,
                "legacy_extra_dirs": [
                    {"path": p, "exists": Path(p).exists()} for p in extra_legacy_dirs
                ],
            },
            indent=2,
        )
    except Exception as e:
        return json.dumps({"error": f"Failed to inspect memory storage: {str(e)}"})


@mcp.tool(
    name="migrate_legacy_memories",
    description="Migrate memories from legacy src/.modular_kingdom storage into the active memory store."
)
def migrate_legacy_memories(
    dry_run: bool = Field(default=True, description="Preview only when true; write changes when false")
) -> str:
    if not scoped_memory:
        return json.dumps({"error": "Memory system not initialized"})
    report = {"dry_run": dry_run, "migrated": 0, "skipped_duplicates": 0, "sources": []}
    try:
        legacy_roots = _legacy_project_roots()
        for root in legacy_roots:
            source_entry = {"project_root": root, "scopes": {}}
            try:
                src_mgr = _scoped_manager_for_root(root, force_project_local=True)
                for scope in MemoryScope:
                    src_items = src_mgr.list_all(scope)
                    dst_items = scoped_memory.list_all(scope)
                    existing_content = {
                        (x.get("content") or "").strip()
                        for x in dst_items
                        if isinstance(x, dict) and (x.get("content") or "").strip()
                    }
                    moved = 0
                    skipped = 0
                    for item in src_items:
                        if not isinstance(item, dict):
                            continue
                        content = (item.get("content") or "").strip()
                        if not content:
                            continue
                        if content in existing_content:
                            skipped += 1
                            continue
                        if not dry_run:
                            scoped_memory.save(content, scope=scope)
                            existing_content.add(content)
                        moved += 1
                    source_entry["scopes"][scope.value] = {
                        "source_count": len(src_items),
                        "to_migrate": moved,
                        "duplicates": skipped,
                    }
                    report["migrated"] += moved
                    report["skipped_duplicates"] += skipped
            except Exception as e:
                source_entry["error"] = str(e)
            report["sources"].append(source_entry)
        if not legacy_roots:
            report["note"] = "No legacy src/.modular_kingdom storage found."
        return json.dumps(report, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Legacy migration failed: {str(e)}"})


@mcp.tool(
    name="cleanup_legacy_memory_paths",
    description="Archive legacy memory directories under .modular_kingdom/legacy_backups after migration."
)
def cleanup_legacy_memory_paths(
    confirm: bool = Field(default=False, description="Set true to perform archive move")
) -> str:
    try:
        candidates = [Path(p) for p in _legacy_extra_dirs() if Path(p).exists()]
        if not candidates:
            return json.dumps({"status": "nothing_to_cleanup", "paths": []}, indent=2)

        sizes = {}
        for p in candidates:
            total = 0
            for fp in p.rglob("*"):
                if fp.is_file():
                    try:
                        total += fp.stat().st_size
                    except Exception:
                        pass
            sizes[str(p)] = total

        if not confirm:
            return json.dumps(
                {
                    "status": "preview",
                    "paths": [str(p) for p in candidates],
                    "bytes": sizes,
                    "hint": "Call with confirm=true to archive these paths.",
                },
                indent=2,
            )

        backup_root = Path(repo_root) / ".modular_kingdom" / "legacy_backups" / datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        backup_root.mkdir(parents=True, exist_ok=True)
        moved = []
        for p in candidates:
            target = backup_root / p.name
            suffix = 1
            while target.exists():
                target = backup_root / f"{p.name}_{suffix}"
                suffix += 1
            shutil.move(str(p), str(target))
            moved.append({"from": str(p), "to": str(target)})

        return json.dumps(
            {"status": "archived", "backup_root": str(backup_root), "moved": moved},
            indent=2,
        )
    except Exception as e:
        return json.dumps({"error": f"Cleanup failed: {str(e)}"})


if _EXPOSE_EXTRA_TOOLS:
    from tools.code_exec import run_code

    @mcp.tool(
        name="code_execute",
        description="Execute Python code in a sandboxed subprocess and return stdout/stderr"
    )
    def code_execute(
        code: str = Field(description="Python code to execute"),
        timeout_seconds: int = Field(default=15, description="Execution timeout in seconds")
    ) -> str:
        try:
            return _call_tool_safely(run_code, code=code, timeout_seconds=timeout_seconds)
        except Exception as e:
            return json.dumps({"status": "error", "error": str(e)})

    @mcp.tool(
        name="shell_execute",
        description="Run a bounded local shell command in AMK_WORKSPACE_ROOT and return stdout/stderr/cwd. Use for explicit terminal-style tasks."
    )
    def shell_execute(
        command: str = Field(description="Shell command to run"),
        timeout_seconds: int = Field(default=15, description="Execution timeout in seconds")
    ) -> str:
        if not isinstance(command, str) or not command.strip():
            return json.dumps({"status": "error", "error": "Empty command"})
        try:
            proc = subprocess.run(
                command,
                cwd=workspace_root,
                shell=True,
                executable="/bin/bash",
                capture_output=True,
                text=True,
                timeout=timeout_seconds,
            )
            return json.dumps(
                {
                    "status": "success" if proc.returncode == 0 else "error",
                    "returncode": proc.returncode,
                    "stdout": proc.stdout,
                    "stderr": proc.stderr,
                    "cwd": workspace_root,
                }
            )
        except subprocess.TimeoutExpired:
            return json.dumps({"status": "error", "error": f"Command timed out after {timeout_seconds}s"})
        except Exception as e:
            return json.dumps({"status": "error", "error": str(e)})

    @mcp.tool(
        name="text_to_speech",
        description="Convert text to speech using various TTS engines. Can play audio directly or save to file."
    )
    def tts_tool(
        text: str = Field(description="Text to convert to speech"),
        engine: str = Field(default="pyttsx3", description="TTS engine to use: pyttsx3, gtts, or kokoro"),
        voice: str = Field(default="", description="Voice name/ID to use (engine-specific)"),
        speed: float = Field(default=200, description="Speech rate for pyttsx3"),
        play_audio: bool = Field(default=True, description="Whether to play audio immediately")
    ) -> str:
        """Convert text to speech and optionally play it."""
        try:
            kwargs = {
                "engine": engine,
                "speed": speed,
                "play_audio": play_audio
            }
            if voice:
                kwargs["voice"] = voice

            return _call_in_subprocess("tools.tts", "text_to_speech", {"text": text, **kwargs}, timeout=60)
        except Exception as e:
            return json.dumps({"success": False, "error": str(e)})

    @mcp.tool(
        name="speech_to_text",
        description="Convert speech to text using microphone recording or existing audio file"
    )
    def stt_tool(
        duration: int = Field(default=5, description="Recording duration in seconds (ignored if file_path provided)"),
        engine: str = Field(default="whisper", description="STT engine: whisper or speech_recognition"),
        model: str = Field(default="base", description="Model size for Whisper: tiny, base, small, medium, large"),
        language: str = Field(default="", description="Language code (e.g., en, es, fr)"),
        file_path: str = Field(default="", description="Path to existing audio file to transcribe")
    ) -> str:
        """Record audio and convert it to text, or transcribe an existing audio file."""
        try:
            kwargs = {
                "duration": duration,
                "engine": engine,
                "model": model
            }
            if language:
                kwargs["language"] = language
            if file_path:
                kwargs["file_path"] = file_path

            return _call_in_subprocess("tools.stt", "speech_to_text", kwargs, timeout=max(60, duration + 45))
        except Exception as e:
            return json.dumps({"success": False, "error": str(e)})

@mcp.prompt(
    name="local_file_answer",
    title="Answer From Local Files",
    description="Use when the user mentioned local files with @ or by exact filename."
)
def local_file_answer(user_request: str, document_context: str = "") -> list[dict]:
    return [
        {
            "role": "user",
            "content": (
                "Answer the user's request from the provided local file context. "
                "Do not claim the file is missing if document_context is non-empty.\n\n"
                f"User request:\n{user_request}\n\n"
                f"DOCUMENTS:\n{document_context}"
            ),
        }
    ]


@mcp.prompt(
    name="juliette_mode_select",
    title="Juliette Mode Selection",
    description="Start Juliette by asking for the mode before persona behavior."
)
def juliette_mode_select() -> list[dict]:
    return [
        {
            "role": "user",
            "content": (
                "Juliette wake-up: read ./Juliette/protocol.md first when available, "
                "then ask which mode the user wants: Learn, Intimate, or Anchor. "
                "Do not invent missing local files."
            ),
        }
    ]


@mcp.prompt(
    name="local_first_agent",
    title="Local-First Agent",
    description="Default behavior for AMK/July: local files and memory before web."
)
def local_first_agent(task: str = "") -> list[dict]:
    return [
        {
            "role": "user",
            "content": (
                "Work local-first. Treat AMK_WORKSPACE_ROOT as the user's cwd. "
                "Use exact local files when named, use memory for durable context, "
                "and use web only when local sources are insufficient or current information is needed.\n\n"
                f"Task:\n{task}"
            ),
        }
    ]

if _EXPOSE_RESOURCES:
    # --- RESOURCES for @ functionality ---
    @mcp.resource("docs://documents", mime_type="application/json")
    def list_documents() -> str:
        """Returns workspace document paths for @ mentions."""
        try:
            doc_ids = []
            for root, dirs, files in os.walk(workspace_root):
                dirs[:] = [
                    d for d in dirs
                    if d not in _DOC_SKIP_DIRS and not d.startswith(".")
                ]
                for name in files:
                    if name.startswith("."):
                        continue
                    suffix = os.path.splitext(name)[1].lower()
                    if suffix not in _DOC_EXTENSIONS:
                        continue
                    path = os.path.join(root, name)
                    try:
                        if os.path.getsize(path) > 2_000_000:
                            continue
                    except OSError:
                        continue
                    doc_ids.append(os.path.relpath(path, workspace_root))
            return json.dumps(sorted(dict.fromkeys(doc_ids))[:500])
        except Exception:
            return json.dumps([])

    @mcp.resource("docs://documents/{doc_id}", mime_type="text/plain")
    def get_document_content(doc_id: str) -> str:
        """Returns the content of a specific document."""
        try:
            file_path = _resolve_target_path(doc_id)

            if not file_path.exists():
                raise ValueError(f"Document {doc_id} not found")

            if doc_id.lower().endswith(".pdf"):
                import fitz

                doc = fitz.open(str(file_path))
                text = "".join(page.get_text() for page in doc)
                doc.close()
                return text

            return file_path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return f"Error: Could not retrieve content for {doc_id}: {str(e)}"

def _warm_rag_v1_background() -> None:
    # Important: never print to stdout; MCP stdio transport uses stdout for protocol messages.
    try:
        from rag import fetch_1 as _rag_v1
        _ = _rag_v1.get_rag_pipeline()
        _elog("[HOST] RAG v1 warmed (background)\n")
    except Exception as _e:
        _elog(f"[HOST] RAG warm skipped/failed: {_e}\n")
    finally:
        pass


# Do not block MCP initialization on model loads/indexing.
# Default is OFF because warm-up can be slow and can also require network (e.g. vector DB).
if os.environ.get("MCP_WARM_RAG_ON_START", "0").lower() in ("1", "true", "yes", "y", "on"):
    try:
        import threading

        threading.Thread(target=_warm_rag_v1_background, daemon=True).start()
    except Exception as _e:
        _elog(f"[HOST] RAG warm thread failed: {_e}\n")

_elog("[HOST] Entering MCP stdio event loop\n")

import asyncio
from importlib.metadata import version as _pkg_version

import mcp.types as _mcp_types
from mcp.shared.version import SUPPORTED_PROTOCOL_VERSIONS as _SUPPORTED_PROTOCOL_VERSIONS

# This server does NOT use the SDK's stdio transport because stdin async-IO is
# unreliable in this environment. We run a tiny synchronous JSON-RPC loop and
# delegate the actual tool implementations to FastMCP.
_server_info = _mcp_types.Implementation(
    name="unified_knowledge_agent_host",
    version=_pkg_version("mcp"),
)
_server_capabilities = _mcp_types.ServerCapabilities(
    tools=_mcp_types.ToolsCapability(listChanged=False),
    resources=_mcp_types.ResourcesCapability(subscribe=False, listChanged=False),
    prompts=_mcp_types.PromptsCapability(listChanged=False),
)

_initialized = False

_use_content_length_framing = False


def _write_jsonrpc(message: _mcp_types.JSONRPCMessage) -> None:
    payload = message.model_dump_json(by_alias=True, exclude_none=True)
    if _use_content_length_framing:
        data = payload.encode("utf-8")
        header = f"Content-Length: {len(data)}\r\n\r\n".encode("ascii")
        sys.stdout.buffer.write(header)
        sys.stdout.buffer.write(data)
        sys.stdout.buffer.flush()
    else:
        sys.stdout.write(payload + "\n")
        sys.stdout.flush()


def _read_jsonrpc_payload() -> str | None:
    """
    Support both:
    - newline-delimited JSON (legacy stdio clients)
    - Content-Length framed JSON (rmcp / LSP-style transport)
    """
    global _use_content_length_framing

    first = sys.stdin.buffer.readline()
    if not first:
        return None
    if _DEBUG_PROTOCOL:
        _elog(f"[HOST][proto] first_line={first[:120]!r}\n")

    lower = first.lower()
    if lower.startswith(b"content-length:"):
        _use_content_length_framing = True

        header_lines = [first]
        while True:
            line = sys.stdin.buffer.readline()
            if not line:
                return None
            header_lines.append(line)
            if line in (b"\r\n", b"\n"):
                break
        if _DEBUG_PROTOCOL:
            _elog(f"[HOST][proto] header_lines={len(header_lines)}\n")

        content_length = None
        for raw in header_lines:
            try:
                text = raw.decode("ascii", errors="ignore").strip()
            except Exception:
                continue
            if not text:
                continue
            if ":" not in text:
                continue
            k, v = text.split(":", 1)
            if k.strip().lower() == "content-length":
                try:
                    content_length = int(v.strip())
                except ValueError:
                    content_length = None

        if content_length is None or content_length < 0:
            if _DEBUG_PROTOCOL:
                _elog("[HOST][proto] missing/invalid content-length\n")
            return None

        remaining = content_length
        chunks: list[bytes] = []
        while remaining > 0:
            chunk = sys.stdin.buffer.read(remaining)
            if not chunk:
                return None
            chunks.append(chunk)
            remaining -= len(chunk)
        body = b"".join(chunks)
        return body.decode("utf-8", errors="replace")

    # Newline-delimited JSON
    return first.decode("utf-8", errors="replace")


def _send_error(req_id: _mcp_types.RequestId, code: int, message: str) -> None:
    err = _mcp_types.JSONRPCError(
        jsonrpc="2.0",
        id=req_id,
        error=_mcp_types.ErrorData(code=code, message=message),
    )
    _write_jsonrpc(_mcp_types.JSONRPCMessage(err))


def _send_result(req_id: _mcp_types.RequestId, result: _mcp_types.ServerResult) -> None:
    resp = _mcp_types.JSONRPCResponse(
        jsonrpc="2.0",
        id=req_id,
        result=result.model_dump(by_alias=True, mode="json", exclude_none=True),
    )
    _write_jsonrpc(_mcp_types.JSONRPCMessage(resp))


_loop = asyncio.new_event_loop()
asyncio.set_event_loop(_loop)

try:
    while True:
        payload = _read_jsonrpc_payload()
        if payload is None:
            break

        try:
            msg = _mcp_types.JSONRPCMessage.model_validate_json(payload)
        except Exception:
            # Can't respond without a request id.
            if _DEBUG_PROTOCOL:
                _elog(f"[HOST][proto] failed to parse jsonrpc payload (len={len(payload) if payload else 0})\n")
            continue

        if not isinstance(msg.root, _mcp_types.JSONRPCRequest):
            # Notifications/responses from the client are ignored (Codex sends initialized notification).
            continue

        req_id = msg.root.id
        if _DEBUG_PROTOCOL:
            _elog(f"[HOST][proto] request method={getattr(msg.root,'method',None)!r} id={req_id!r}\n")
        try:
            req = _mcp_types.ClientRequest.model_validate(
                msg.root.model_dump(by_alias=True, mode="json", exclude_none=True)
            )
        except Exception:
            _send_error(req_id, _mcp_types.INVALID_PARAMS, "Invalid request parameters")
            continue

        try:
            match req.root:
                case _mcp_types.InitializeRequest(params=params):
                    requested = params.protocolVersion
                    protocol = (
                        requested if requested in _SUPPORTED_PROTOCOL_VERSIONS else _mcp_types.LATEST_PROTOCOL_VERSION
                    )
                    init = _mcp_types.InitializeResult(
                        protocolVersion=protocol,
                        capabilities=_server_capabilities,
                        serverInfo=_server_info,
                        instructions=None,
                    )
                    _send_result(req_id, _mcp_types.ServerResult(init))
                    _initialized = True

                case _mcp_types.PingRequest():
                    _send_result(req_id, _mcp_types.ServerResult(_mcp_types.EmptyResult()))

                case _mcp_types.ListToolsRequest():
                    if not _initialized:
                        raise RuntimeError("Received tools/list before initialization")
                    tools = _loop.run_until_complete(mcp.list_tools())
                    _send_result(req_id, _mcp_types.ServerResult(_mcp_types.ListToolsResult(tools=tools)))

                case _mcp_types.CallToolRequest(params=params):
                    if not _initialized:
                        raise RuntimeError("Received tools/call before initialization")
                    args = params.arguments or {}
                    out = _loop.run_until_complete(mcp.call_tool(params.name, args))
                    # FastMCP returns:
                    # - CallToolResult (already in protocol form), or
                    # - list[ContentBlock] for unstructured output, or
                    # - tuple(list[ContentBlock], dict) for structured output, or
                    # - dict for structured output.
                    if isinstance(out, _mcp_types.CallToolResult):
                        result = out
                    elif isinstance(out, tuple) and len(out) == 2:
                        content, structured = out
                        result = _mcp_types.CallToolResult(
                            content=list(content),
                            structuredContent=structured,
                            isError=False,
                        )
                    elif isinstance(out, dict):
                        result = _mcp_types.CallToolResult(content=[], structuredContent=out, isError=False)
                    else:
                        result = _mcp_types.CallToolResult(content=list(out), structuredContent=None, isError=False)
                    _send_result(req_id, _mcp_types.ServerResult(result))

                case _mcp_types.ListResourcesRequest():
                    resources = _loop.run_until_complete(mcp.list_resources())
                    _send_result(req_id, _mcp_types.ServerResult(_mcp_types.ListResourcesResult(resources=resources)))

                case _mcp_types.ReadResourceRequest(params=params):
                    contents = _loop.run_until_complete(mcp.read_resource(params.uri))
                    _send_result(req_id, _mcp_types.ServerResult(_mcp_types.ReadResourceResult(contents=contents)))

                case _mcp_types.ListPromptsRequest():
                    prompts = _loop.run_until_complete(mcp.list_prompts())
                    _send_result(req_id, _mcp_types.ServerResult(_mcp_types.ListPromptsResult(prompts=prompts)))

                case _mcp_types.GetPromptRequest(params=params):
                    prompt = _loop.run_until_complete(mcp.get_prompt(params.name, params.arguments))
                    _send_result(req_id, _mcp_types.ServerResult(prompt))

                case _mcp_types.ListResourceTemplatesRequest():
                    templates = _loop.run_until_complete(mcp.list_resource_templates())
                    _send_result(
                        req_id, _mcp_types.ServerResult(_mcp_types.ListResourceTemplatesResult(resourceTemplates=templates))
                    )

                case _:
                    _send_error(req_id, _mcp_types.METHOD_NOT_FOUND, f"Unsupported method: {msg.root.method}")
        except Exception as e:
            _send_error(req_id, _mcp_types.INTERNAL_ERROR, str(e))
finally:
    try:
        _loop.close()
    except Exception:
        pass
